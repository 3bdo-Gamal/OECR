from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def is_arabic(text: str) -> bool:
    """Check if any character in the text is Arabic."""
    for char in text:
        if '\u0600' <= char <= '\u06FF':
            return True
    return False

def export_text_to_docx(text: str, output_path: str) -> bool:
    try:
        doc = Document()
        title = doc.add_heading('OECR Extracted Document', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        lines = text.split('\n')
        for line in lines:
            paragraph = doc.add_paragraph()
            
            if is_arabic(line):
                # Set paragraph direction to RTL for Arabic text
                pPr = paragraph._element.get_or_add_pPr()
                from docx.oxml import OxmlElement
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
            run = paragraph.add_run(line)
            font = run.font
            font.name = 'Arial'
            font.size = Pt(12)
            
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error saving to DOCX: {e}")
        return False
